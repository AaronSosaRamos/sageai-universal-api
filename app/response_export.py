"""
Exportación de respuestas (Markdown) a Word (.docx) y PDF.
Usado por el endpoint /export/response.
"""

from __future__ import annotations

import io
import re
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fpdf import FPDF


MAX_EXPORT_CHARS = 500_000

# Posibles rutas de fuente DejaVu (p. ej. Docker: fonts-dejavu-core)
DEJAVU_CANDIDATES = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    "/usr/local/share/fonts/DejaVuSans.ttf",
)


def _find_dejavu_regular() -> Optional[str]:
    import os

    for path in DEJAVU_CANDIDATES:
        if "Bold" in path:
            continue
        if os.path.isfile(path):
            return path
    return None


def _find_dejavu_bold() -> Optional[str]:
    import os

    p = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
    if os.path.isfile(p):
        return p
    return _find_dejavu_regular()


def _sanitize_title(title: str, max_len: int = 80) -> str:
    t = re.sub(r'[<>:"/\\|?*]', "", title.strip() or "respuesta")
    return t[:max_len] if t else "respuesta"


def _parse_table_block(lines: List[str], start: int) -> Tuple[Optional[List[List[str]]], int]:
    """Detecta tabla GFM (| a | b |) desde start. Devuelve filas y nuevo índice."""
    if start >= len(lines):
        return None, start
    row_lines: List[str] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row_lines.append(lines[i].strip())
        i += 1
    if len(row_lines) < 1:
        return None, start

    def split_row(line: str) -> List[str]:
        inner = line.strip()
        if inner.startswith("|"):
            inner = inner[1:]
        if inner.endswith("|"):
            inner = inner[:-1]
        return [c.strip() for c in inner.split("|")]

    rows: List[List[str]] = []
    for rl in row_lines:
        cells = split_row(rl)
        if cells and all(re.match(r"^[\s:-]+$", c) for c in cells):
            continue
        if cells:
            rows.append(cells)
    if not rows:
        return None, start
    return rows, i


def _add_markdown_to_doc(doc: Document, md: str) -> None:
    lines = md.replace("\r\n", "\n").split("\n")
    i = 0
    in_code = False
    code_buf: List[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table, new_i = _parse_table_block(lines, i)
            if table:
                ncols = max(len(r) for r in table)
                t = doc.add_table(rows=len(table), cols=ncols)
                try:
                    t.style = "Table Grid"
                except (ValueError, KeyError):
                    pass
                for ri, row in enumerate(table):
                    for ci in range(ncols):
                        text = row[ci] if ci < len(row) else ""
                        t.cell(ri, ci).text = text
                i = new_i
                continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=4)
        elif stripped.startswith("##### "):
            doc.add_heading(stripped[6:].strip(), level=5)
        elif stripped.startswith("###### "):
            doc.add_heading(stripped[7:].strip(), level=6)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        elif stripped == "":
            pass
        else:
            doc.add_paragraph(line)

        i += 1


def markdown_to_docx(markdown: str, title: str = "Respuesta") -> bytes:
    if len(markdown) > MAX_EXPORT_CHARS:
        markdown = markdown[:MAX_EXPORT_CHARS] + "\n\n[… contenido truncado …]"

    doc = Document()
    h = doc.add_heading(_sanitize_title(title), 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    _add_markdown_to_doc(doc, markdown)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


class _ExportPDF(FPDF):
    def __init__(self) -> None:
        super().__init__()
        self.set_auto_page_break(auto=True, margin=18)


def markdown_to_pdf(markdown: str, title: str = "Respuesta") -> bytes:
    if len(markdown) > MAX_EXPORT_CHARS:
        markdown = markdown[:MAX_EXPORT_CHARS] + "\n\n[… contenido truncado …]"

    pdf = _ExportPDF()
    pdf.set_margins(14, 14, 14)
    regular = _find_dejavu_regular()
    bold = _find_dejavu_bold()

    if regular:
        pdf.add_font("DejaVu", "", regular, uni=True)
        if bold:
            pdf.add_font("DejaVu", "B", bold, uni=True)
        pdf.set_font("DejaVu", "", 11)
        use_unicode = True
    else:
        pdf.set_font("Helvetica", "", 11)
        use_unicode = False

    pdf.add_page()

    safe_title = _sanitize_title(title)
    if use_unicode:
        pdf.set_font("DejaVu", "B", 14)
        pdf.multi_cell(0, 9, safe_title, align="C")
    else:
        pdf.set_font("Helvetica", "B", 14)
        pdf.multi_cell(0, 9, safe_title.encode("latin-1", "replace").decode("latin-1"), align="C")
    pdf.ln(6)

    if use_unicode:
        pdf.set_font("DejaVu", "", 11)
    else:
        pdf.set_font("Helvetica", "", 11)

    lines = markdown.replace("\r\n", "\n").split("\n")
    i = 0
    in_code = False
    code_buf: List[str] = []
    effective_w = pdf.w - pdf.l_margin - pdf.r_margin

    def out_text(text: str, size: int = 11, is_bold: bool = False) -> None:
        if use_unicode:
            pdf.set_font("DejaVu", "B" if is_bold else "", size)
            pdf.multi_cell(effective_w, size * 0.45, text)
        else:
            pdf.set_font("Helvetica", "B" if is_bold else "", size)
            safe = text.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(effective_w, size * 0.45, safe)
        pdf.ln(2)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                block = "\n".join(code_buf)
                if use_unicode:
                    pdf.set_font("DejaVu", "", 9)
                    pdf.multi_cell(effective_w, 4.5, block)
                else:
                    pdf.set_font("Courier", "", 9)
                    pdf.multi_cell(
                        effective_w,
                        4.5,
                        block.encode("latin-1", "replace").decode("latin-1"),
                    )
                pdf.ln(3)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table, new_i = _parse_table_block(lines, i)
            if table:
                for row in table:
                    row_txt = "  |  ".join(row)
                    out_text(row_txt, 10)
                pdf.ln(2)
                i = new_i
                continue

        if stripped.startswith("# "):
            out_text(stripped[2:].strip(), 16, True)
        elif stripped.startswith("## "):
            out_text(stripped[3:].strip(), 14, True)
        elif stripped.startswith("### "):
            out_text(stripped[4:].strip(), 12, True)
        elif stripped.startswith("#### "):
            out_text(stripped[5:].strip(), 11, True)
        elif stripped.startswith(("- ", "* ")):
            out_text("• " + stripped[2:].strip(), 11)
        elif re.match(r"^\d+\.\s+", stripped):
            out_text(stripped, 11)
        elif stripped == "":
            pdf.ln(3)
        else:
            out_text(stripped, 11)

        i += 1

    return pdf.output()


def export_markdown(markdown: str, fmt: str, title: Optional[str] = None) -> Tuple[bytes, str, str]:
    """
    Retorna (bytes, media_type, filename).
    fmt: 'docx' | 'pdf'
    """
    t = title or "respuesta-sageai"
    safe = _sanitize_title(t)
    if fmt == "docx":
        raw = markdown_to_docx(markdown, safe)
        data = bytes(raw) if isinstance(raw, (bytearray, memoryview)) else raw
        return (
            data,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{safe}.docx",
        )
    if fmt == "pdf":
        raw = markdown_to_pdf(markdown, safe)
        data = bytes(raw) if isinstance(raw, (bytearray, memoryview)) else raw
        return (
            data,
            "application/pdf",
            f"{safe}.pdf",
        )
    raise ValueError(f"Formato no soportado: {fmt}")
